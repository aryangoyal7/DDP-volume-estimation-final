#!/usr/bin/env python3
"""Research / thesis style DOCX report on the calibrated 2D volume
estimation framework. Mirrors the numerical results of latex_report/main.tex
and is restricted to the IITB three-pile evaluation.
"""
from pathlib import Path
import json
import csv

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path("/home/bheeshmsharma/Karthikeyan_new/a_g/VLA-volume-estimation")
FIGS = ROOT / "report_docx" / "figures"
FF   = ROOT / "eval_results" / "`"
OUT_DOCX = ROOT / "report_docx" / "VolumeEstimation_Report.docx"

# ── data --------------------------------------------------------------------
summary = json.loads((FF / "calibration_summary.json").read_text())
iitb_gt = json.loads((ROOT / "eval_results" / "iitb_gt.json").read_text())

# Numbers used in the body — kept consistent with latex_report/main.tex
PAPER = {
    "pile1": dict(mass_g=670, gt_L=18.0, raw_mean=13.4, raw_std=2.1, cv_pct=15.6, cf=1.38, cal_L=18.5, err_pct=2.8),
    "pile2": dict(mass_g=480, gt_L=18.0, raw_mean=37.8, raw_std=17.7, cv_pct=46.9, cf=0.48, cal_L=18.1, err_pct=0.6),
    "pile3": dict(mass_g=660, gt_L=18.0, raw_mean=19.0, raw_std=5.4,  cv_pct=28.5, cf=1.02, cal_L=19.4, err_pct=7.8),
}
GLOBAL_CF_MEDIAN = 0.86

# ── doc helpers -------------------------------------------------------------
doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)


def H(level, text):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def P(text, italic=False, bold=False, justify=True, size=11):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    return p


def caption(fig_no, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(f"Figure {fig_no}. "); r.bold = True
    r.font.name = "Times New Roman"; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.italic = True
    r2.font.name = "Times New Roman"; r2.font.size = Pt(10)


def add_figure(filename, fig_no, cap, width_in=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(str(FIGS / filename), width=Inches(width_in))
    caption(fig_no, cap)


def add_table(headers, data, col_widths=None, table_no=None, table_caption=None):
    if table_no is not None and table_caption is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"Table {table_no}. ")
        r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
        r2 = p.add_run(table_caption); r2.italic = True
        r2.font.name = "Times New Roman"; r2.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = ""
            para = row[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)


# =============================================================================
# Title
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Calibrated Monocular Depth for Solid Waste\n"
                  "Volume Estimation from Overhead Imagery")
r.bold = True; r.font.size = Pt(18); r.font.name = "Times New Roman"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A fixed-camera, fixed-box pipeline built on DepthAnythingV2-Metric\n"
                "with a one-parameter scene-specific calibration factor")
r.italic = True; r.font.size = Pt(12); r.font.name = "Times New Roman"
doc.add_paragraph()

# =============================================================================
# Abstract
# =============================================================================
H(1, "Abstract")
P("Accurate volume estimation of solid waste piles is a practical "
  "requirement at material recovery facilities, transfer stations and "
  "audit sites, yet conventional approaches rely on either manual "
  "measurement or expensive 3D sensors. This work investigates a "
  "low-cost alternative: a single overhead RGB image captured from a "
  "fixed tripod, combined with a monocular metric depth model, to "
  "estimate pile volume. The scene is a 60 cm by 60 cm collection frame "
  "placed on a flat surface; the camera is mounted 65 cm above the frame "
  "floor and the frame fills almost the entire image. Depth maps are "
  "produced by DepthAnythingV2-Metric-Indoor-Large and converted to "
  "height maps through tilted plane fitting on the deepest 1 percent of "
  "valid pixels. A scene-specific calibration factor, derived from a "
  "small set of ground-truth labelled captures, corrects for systematic "
  "depth compression and material fluffing. Evaluated on 34 images "
  "across three mixed-composition piles at the IIT Bombay site (670 g, "
  "480 g and 660 g net waste mass, all packed to a geometric envelope of "
  "18 L), the calibrated system achieves a volumetric error of 0.6 to "
  "7.8 percent across the three conditions, with 100 percent depth "
  "sanity verification on every capture. These results demonstrate that "
  "monocular depth, combined with a simple one-parameter calibration, "
  "is a viable pathway toward practical in-situ waste volume "
  "estimation, and that the reliability of the calibration improves as "
  "more images per pile are aggregated.")

# =============================================================================
# 1 Introduction
# =============================================================================
H(1, "1. Introduction")
P("Solid waste management requires reliable volume and mass numbers at "
  "almost every stage of the value chain. Field audits, informal "
  "collection points and material recovery facilities all benefit from "
  "quick, low-cost in-situ volumetric estimates of small piles. The de "
  "facto options are either weighing, which gives no volume information, "
  "or full 3D capture (lidar, structured-light, RGBD), neither of which "
  "is practical at scale in cluttered, dusty, outdoor environments.")
P("The central idea explored here is to combine a fixed overhead camera "
  "with a monocular metric depth network. Given a known scene footprint "
  "(the floor area enclosed by the containment frame) and an estimate "
  "of each pixel's height above the floor, volume is simply the "
  "integral of the height map over that footprint. The challenge is "
  "that monocular depth models are calibrated to indoor scenes at "
  "typical human-viewing distances; at very close range (65 cm "
  "camera-to-floor), the relative depth span of a 20 cm pile is "
  "compressed to a small fraction of the model's output range. A "
  "per-scene calibration factor bridges this gap.")
P("Specifically, this report contributes:")
bullet("A complete pipeline, from acquisition geometry to final volume "
       "number, that runs on a single RGB image of a constrained scene.")
bullet("A one-parameter calibration scheme: a scalar that corrects the "
       "residual systematic bias of the monocular depth model after a "
       "plane-fit metric scaling.")
bullet("An empirical evaluation on a 34-image IIT Bombay dataset of three "
       "controlled waste piles, showing post-calibration errors of 0.6 "
       "to 7.8 percent and full depth sanity on every capture.")
bullet("An analysis of how the calibration factor stabilises and the "
       "variance of the estimate shrinks as more labelled images are "
       "aggregated, motivating the data-collection protocol for new "
       "deployments.")

# =============================================================================
# 2 Related Work
# =============================================================================
H(1, "2. Related Work")
P("Single-image volume estimation has a long history in agriculture and "
  "industry. Classical photogrammetric approaches rely on multi-view "
  "stereo or structure-from-motion to triangulate dense 3D points and "
  "integrate them under a fitted ground plane. These methods are "
  "accurate but require either careful calibration patterns or many "
  "views per pile, neither of which is convenient for routine in-situ "
  "waste audits.")
P("Monocular depth estimation has progressed substantially. "
  "DepthAnythingV2 (Yang et al., 2024) offers a metric-trained "
  "transformer that predicts depth in metres from a single RGB image, "
  "and an Indoor-Large variant fine-tuned for short-range indoor "
  "geometry. The model is small enough to run on a single GPU and is "
  "the depth backbone used throughout this report. Vision-language "
  "segmentation models such as Segment-Anything (Kirillov et al., 2023) "
  "can isolate the pile from background, but for the controlled "
  "fixed-frame setup used here, the frame itself fills the field of "
  "view and no segmentation is needed.")

# =============================================================================
# 3 Dataset
# =============================================================================
H(1, "3. Dataset")
P("All experiments use the IIT Bombay waste-pile dataset. Three piles of "
  "mixed solid waste (plastic, paper, cardboard, metal cans, single-use "
  "wrappers) were prepared inside a white-chalked containment frame "
  "with an interior footprint of 60 cm by 60 cm (area 0.36 m^2). An "
  "iPhone was mounted on a tripod 65 cm above the frame floor, oriented "
  "vertically downward. At this height the frame fills the vast "
  "majority of the image, providing a natural scene boundary and a "
  "stable reference geometry. Figure 1 shows the collection site and "
  "representative captures.")
add_figure("fig9_site.png", 1,
           "IIT Bombay solid waste collection site. Top row: site "
           "overview, showing the tripod and the 60 by 60 cm white frame. "
           "Bottom row: representative overhead captures of the three "
           "pile conditions.",
           width_in=6.5)
P("Three mass conditions of mixed solid waste were studied. Net masses "
  "below are gross minus a 270 g tare:")
add_table(
    ["Pile", "Material", "Net mass (g)", "Geometric GT volume (L)",
     "Number of images"],
    [
        ["pile1 (A)", "Mixed plastic, paper, cardboard, cans, wrappers",
         "670", "18.0", "8"],
        ["pile2 (B)", "Mixed plastic, paper, cardboard, cans, wrappers",
         "480", "18.0", "13"],
        ["pile3 (C)", "Mixed plastic, paper, cardboard, cans, wrappers",
         "660", "18.0", "13"],
        ["Total",     "",                                            "1810 g",
         "",                                                          "34"],
    ],
    col_widths=[1.0, 3.0, 1.0, 1.6, 1.2],
    table_no=1,
    table_caption="IIT Bombay three-pile dataset. The geometric GT volume "
                  "of 18 L corresponds to a settled bounding box of "
                  "20 cm by 30 cm by 30 cm. Bulk volume may exceed this "
                  "due to material fluffing.",
)
P("For each pile condition, the camera tripod was repositioned within a "
  "small radius around the vertical axis and captures were taken across "
  "different pile arrangements and ambient lighting conditions. This "
  "multi-view protocol provides a distribution of raw depth-based "
  "volume estimates for each pile, from which a robust calibration "
  "factor can be derived. The captured images are divided into a "
  "calibration partition (used to derive the calibration factor) and "
  "an evaluation partition (held out for blind evaluation), simulating "
  "deployment on unseen captures from the same scene type.")

# =============================================================================
# 4 Methodology
# =============================================================================
H(1, "4. Methodology")
P("The volume estimation pipeline consists of five stages: depth "
  "inference, floor-plane fitting, height-map computation, full-frame "
  "volume integration, and per-condition calibration. Figure 2 "
  "summarises the framework and Figure 3 shows the acquisition "
  "geometry schematically.")
add_figure("fig1_framework.png", 2,
           "Volume estimation framework. The top row recovers a depth "
           "map and fits a planar floor reference. The middle row "
           "converts the depth map into a metric height field and "
           "integrates volume over the full frame. The bottom row "
           "applies an empirical calibration factor estimated from a "
           "small labelled set of piles.",
           width_in=6.6)
add_figure("fig2_geometry.png", 3,
           "Acquisition geometry. The iPhone is mounted at H_cam = "
           "0.65 m above a 60 by 60 cm white frame. The frame interior "
           "fills the image, so its corners and edges serve as the "
           "floor reference for plane fitting and metric scaling.",
           width_in=4.5)

H(2, "4.1 Depth inference")
P("The depth model is DepthAnythingV2-Metric-Indoor-Large, a "
  "transformer-based monocular depth estimator fine-tuned for metric "
  "depth prediction in indoor environments. Unlike relative-depth "
  "models, it produces depth values in metres, which enables direct "
  "conversion to physical heights without additional scale recovery. "
  "Inference is performed on a single GPU with mixed-precision "
  "computation. At the operating distance of 65 cm the model typically "
  "outputs floor depths in the range 1.75 m to 2.25 m. This systematic "
  "offset is handled in the floor-fitting stage.")

H(2, "4.2 Floor-plane fitting and metric scaling")
P("The floor of the containment frame acts as the depth reference. "
  "Because the camera can have a slight non-vertical tilt, a planar "
  "floor model is fitted rather than assuming a constant floor depth:")
bullet("Floor candidate pixels are the deepest 1 percent of valid depth "
       "values. At 65 cm camera height with the frame filling the image, "
       "these pixels correspond to the frame edges and corners where no "
       "pile material is present.")
bullet("A tilted plane d_plane(x, y) = a x + b y + c is fitted to the "
       "(x, y, d) coordinates of the floor candidates by least squares.")
bullet("A depth-to-height scale factor s is derived by requiring that "
       "the median floor depth corresponds to the known physical camera "
       "height: s = H / median(d_floor) with H = 0.65 m.")

H(2, "4.3 Height map and depth sanity check")
P("The pile height above the local floor at each pixel is")
P("    h(i, j) = (d_plane(i, j) - d(i, j)) * s,",
  italic=True, justify=False)
P("with pixels below the local floor clipped to zero. A fixed minimum "
  "height threshold of 1 cm is applied to suppress noise in regions "
  "with no pile material:")
P("    h_clean(i, j) = max(h(i, j) - 0.01, 0).",
  italic=True, justify=False)
P("Before accepting a height map, a geometric sanity check is applied: "
  "the mean depth in the central 70 percent of the image must be "
  "smaller (closer to the camera) than the mean depth of the 15 percent "
  "border strip. This verifies that the model has correctly recognised "
  "the pile as elevated above the floor. In all 34 captures evaluated "
  "in this study this sanity check passes, with mean centre-to-edge "
  "depth differences of 6.5 cm, 6.9 cm and 4.1 cm for piles 1 through "
  "3 respectively (in scaled physical units).")

H(2, "4.4 Volume integration")
P("Because the frame fills the image, the volume is integrated over the "
  "full frame with no segmentation step. The pixel area in world units "
  "is A_px = A_box / N, where A_box = 0.36 m^2 and N is the total number "
  "of image pixels. The raw volume is")
P("    V_raw = sum_{i, j} h_clean(i, j) * A_px,",
  italic=True, justify=False)
P("converted from cubic metres to litres by multiplying by 1000.")

H(2, "4.5 Calibration factor")
P("V_raw systematically underestimates or overestimates the true pile "
  "volume for two reasons:")
bullet("Depth compression. At very close range the model's relative "
       "depth span is compressed relative to the physical height "
       "difference. The metric scale factor s partially corrects for "
       "this, but a residual bias remains.")
bullet("Material fluffing. The geometric GT volume (18 L) is measured "
       "on a settled bounding box, but the actual bulk volume of the "
       "arranged pile can differ due to inter-piece air pockets.")
P("To absorb both effects with a single correction, a per-condition "
  "calibration factor is computed from the calibration partition:")
P("    CF = V_GT / mean(V_raw on calibration partition).",
  italic=True, justify=False)
P("At inference time the final estimate is V_est = CF * V_raw.")

# =============================================================================
# 5 Experimental setup
# =============================================================================
H(1, "5. Experimental Setup")
P("All experiments were run on a single CUDA GPU with mixed-precision "
  "inference. DepthAnythingV2 was loaded from the "
  "depth-anything/Depth-Anything-V2-Metric-Indoor-Large-hf Hugging Face "
  "checkpoint. The plane-fit and volume-integration code is implemented "
  "in NumPy on the CPU. The 34 IITB images were processed in a single "
  "batch.")
add_table(
    ["Hyperparameter", "Value", "Notes"],
    [
        ["Depth model",                "DepthAnythingV2-Metric-Indoor-Large", "Hugging Face HF checkpoint"],
        ["Camera height H_cam",        "0.65 m",                              "tripod-mounted iPhone"],
        ["Box interior area A_box",    "0.36 m^2",                            "60 x 60 cm white frame"],
        ["Floor candidate percentile", "99 percent",                          "deepest pixels = floor"],
        ["Height noise threshold",     "0.01 m (1 cm)",                       "fixed, not adaptive"],
        ["Sanity check region",        "centre 70%, border 15%",              "centre depth < border depth"],
        ["GT volume per pile",         "18 L",                                "20 by 30 by 30 cm envelope"],
    ],
    col_widths=[2.4, 2.4, 1.8],
    table_no=2,
    table_caption="Hyperparameters used throughout the IITB three-pile "
                  "evaluation.",
)

# =============================================================================
# 6 Results
# =============================================================================
H(1, "6. Results")

H(2, "6.1 Qualitative outputs and depth sanity")
P("Figure 4 shows the three-panel output produced for a representative "
  "image from each pile condition: the DepthAnythingV2 depth map "
  "(brighter is closer to the camera) and the calibrated height heatmap "
  "blended onto the RGB image (warmer is taller). The depth maps show "
  "a clear central elevation corresponding to the pile, surrounded by "
  "the flat frame floor at the image periphery, confirming that the "
  "model is detecting pile geometry correctly in every case.")
add_figure("fig8_qualitative.png", 4,
           "Qualitative outputs of the pipeline on one image from each "
           "pile. Left: DepthAnythingV2 depth map with the frame visible "
           "as the floor reference. Right: height heatmap overlaid on "
           "the RGB image. All 34 captures pass the centre-vs-edge "
           "depth sanity check.",
           width_in=6.5)

H(2, "6.2 Pile-level volume estimates and calibration factors")
P("Table 3 summarises the raw predicted volume, derived calibration "
  "factor, and final calibrated volume for each pile condition. For "
  "pile1 (670 g net) the model consistently predicts a raw volume of "
  "about 13 L (CV = 15.6 percent) and the calibration factor of 1.38 "
  "corrects this to 18.5 L, within 2.8 percent of GT. For pile2 (480 g) "
  "the raw prediction is much larger (37.8 L, CV = 46.9 percent) and a "
  "calibration factor of 0.48 brings it back to 18.1 L, within 0.6 "
  "percent. For pile3 (660 g) the raw mean of 19.0 L requires only a "
  "calibration factor of 1.02 and the calibrated volume is 19.4 L, an "
  "error of 7.8 percent.")
add_table(
    ["Pile", "Mass (g)", "GT (L)", "Raw mean (L)", "Raw std (L)",
     "CV (%)", "CF", "Calibrated (L)", "Error (%)"],
    [
        ["pile1 (A)", "670", "18.0", "13.4", "2.1",  "15.6", "1.38", "18.5", "2.8"],
        ["pile2 (B)", "480", "18.0", "37.8", "17.7", "46.9", "0.48", "18.1", "0.6"],
        ["pile3 (C)", "660", "18.0", "19.0", "5.4",  "28.5", "1.02", "19.4", "7.8"],
        ["global median CF",
         "", "", "", "", "", f"{GLOBAL_CF_MEDIAN:.2f}", "", ""],
    ],
    col_widths=[0.9, 0.7, 0.6, 0.9, 0.9, 0.7, 0.6, 1.0, 0.8],
    table_no=3,
    table_caption="Volume estimation results per pile. CF is the "
                  "calibration factor derived from the calibration "
                  "partition. Calibrated volume = CF times mean raw "
                  "predicted volume. Error is against the geometric GT "
                  "of 18 L.",
)
P("Figure 5 plots the distribution of raw predicted volumes across all "
  "captures, together with the GT line and the average post-calibration "
  "value (using the global median CF of 0.86). Pile1 sits below the GT "
  "line, pile3 hovers around it, and pile2 sits well above the GT line "
  "with two anomalous captures above 60 L. The per-image scatter is "
  "exactly the kind of bias that a multiplicative CF can absorb on the "
  "pile-level mean.")
add_figure("fig3_per_image_volumes.png", 5,
           "Per-image raw predicted volume for the three IITB piles "
           "using DepthAnythingV2 in full-frame mode. The red dashed "
           "line is the geometric GT of 18 L; the green dotted line is "
           "the average post-calibration volume using the global median "
           "CF of 0.86.",
           width_in=6.6)
P("Figure 6 shows the raw mean, the calibrated mean and the GT side by "
  "side. The calibrated bars are within a few percent of the GT bars for "
  "every pile.")
add_figure("fig4_pile_bars.png", 6,
           "Per-pile raw mean (with one-standard-deviation error bars), "
           "calibrated mean, and ground-truth volume. The calibrated "
           "values match the GT to within 0.6 to 7.8 percent.",
           width_in=6.4)

H(2, "6.3 Stability of the calibration factor with data size")
P("Figure 7 plots the running calibration factor as a function of the "
  "number of images used to compute the running mean. The CF is "
  "volatile after one or two images but converges within five to seven "
  "images for piles 1 and 3, where the per-image predictions are "
  "tightly clustered. Pile2 takes longer to stabilise because of the "
  "two outliers near 60 to 90 L that dominate the running mean until "
  "additional captures dilute them. The horizontal red dashed line is "
  "the global median CF across all three piles.")
add_figure("fig5_cf_running.png", 7,
           "Running calibration factor as a function of the number of "
           "images aggregated. The CF stabilises quickly for tight piles "
           "(pile1, pile3) and more slowly for pile2, which contains "
           "two anomalous high-volume captures. The red dashed line is "
           "the global median CF of 0.86.",
           width_in=6.4)
P("Figure 8 shows the matching story from the variance side. As more "
  "images are aggregated, the coefficient of variation per pile "
  "decreases monotonically for the well-behaved piles. Pile2 stays high "
  "because of the outliers but still narrows once the outliers are "
  "out-numbered by ordinary captures.")
add_figure("fig6_cv_vs_n.png", 8,
           "Coefficient of variation (CV) of the predicted volume as a "
           "function of the number of images used in the running mean. "
           "Tighter piles see a clear downward trend; pile2 remains "
           "high due to outliers.",
           width_in=6.4)

H(2, "6.4 Effect of calibration on the volume error")
P("Figure 9 quantifies the impact of the calibration factor on the "
  "per-pile error. Uncalibrated, the pile2 mean is 110 percent above "
  "GT and the pile1 mean is 26 percent below; pile3 is already close at "
  "6 percent off. After per-pile calibration, all three pile-level "
  "errors fall to single digits (0.6, 2.8 and 7.8 percent). This is the "
  "central practical result of the report: with a per-condition CF, the "
  "raw 26 percent to 110 percent error band shrinks by more than an "
  "order of magnitude.")
add_figure("fig7_calib_effect.png", 9,
           "Effect of the per-pile calibration factor on the volume "
           "error. The uncalibrated bars show the raw absolute "
           "percentage error against GT; the post-CF bars show the "
           "error after applying the per-pile factor from Table 3.",
           width_in=6.0)

H(2, "6.5 Apparent density")
P("Because the GT masses are known per pile, an apparent bulk density "
  "is computed as GT mass divided by predicted volume. Mixed solid "
  "waste typically has a bulk density of 0.10 to 0.30 kg/L. Figure 10 "
  "compares the densities computed against the raw and the calibrated "
  "volumes. The calibrated densities all sit at the lower edge or "
  "below the plausible band, consistent with the fact that the piles "
  "in this study are loose and have a relatively low packing density.")
add_figure("fig10_density.png", 10,
           "Apparent bulk density per pile, computed as GT mass divided "
           "by predicted volume. The shaded green band is the plausible "
           "range for mixed solid waste (0.10 to 0.30 kg/L). Calibrated "
           "densities are close to the plausible range; raw densities "
           "differ markedly per pile because of the uncorrected depth "
           "bias.",
           width_in=6.4)

# =============================================================================
# 7 Discussion
# =============================================================================
H(1, "7. Discussion")
H(2, "7.1 Pile1: consistent under-estimation corrected by CF")
P("For the 670 g pile the model consistently predicts a raw volume of "
  "approximately 13 L with a tight CV of 15.6 percent. A calibration "
  "factor of 1.38 corrects this to 18.5 L, within 2.8 percent of GT. "
  "The under-estimation arises because DepthAnythingV2 is calibrated to "
  "indoor scenes at 1 to 5 m range; at 65 cm the depth variation "
  "across the pile spans a narrow portion of the model's output range, "
  "resulting in compressed height estimates.")
H(2, "7.2 Pile2: high variance and GT uncertainty")
P("Pile2 (480 g) shows the highest raw variance (CV = 46.9 percent) "
  "and a mean raw prediction of 37.8 L, more than twice the GT. Two "
  "captures exhibit anomalously high predicted volumes (up to 87 L), "
  "likely due to camera tilt or specular reflectance from smooth "
  "surfaces in the mixed-waste pile that caused the depth model to "
  "misinterpret near-field material. Excluding those captures, the "
  "remaining pile2 images average around 27 L. The fixed 18 L "
  "geometric GT is also questionable here: a lighter, fluffier pile "
  "can occupy variable bulk volume depending on how it is arranged. "
  "Displacement-based GT measurement is recommended to resolve this "
  "ambiguity.")
H(2, "7.3 Pile3: near-unity CF, good accuracy")
P("For the 660 g pile, the raw mean of 19.0 L requires only a "
  "calibration factor of 1.02, meaning the model already predicts "
  "close to GT without correction. The 7.8 percent post-calibration "
  "error (19.4 L vs 18.0 L) reflects genuine pile-arrangement "
  "variation rather than a model failure.")
H(2, "7.4 Implications for calibration transfer")
P("The per-pile CFs range from 0.48 to 1.38, a factor of nearly three. "
  "This large spread shows that a single global CF is not reliable "
  "across different pile masses or arrangements. The running-CF and "
  "CV analyses (Figures 7 and 8) suggest that five to eight labelled "
  "captures per condition are sufficient to lock in a stable CF for "
  "well-behaved piles. For practical deployment a calibration set "
  "covering the expected range of pile sizes and material types should "
  "be collected for each new site.")
H(2, "7.5 Limitations")
bullet("Fixed camera height assumption. A 10 percent error in H_cam "
       "translates into a 10 percent error in the metric scale before "
       "the CF is applied. The frame still pins the world-area per "
       "pixel, but the height channel still depends on the tripod.")
bullet("Single-scalar bias model. The CF only corrects the mean "
       "behaviour. It cannot fix per-image noise, which is why the CV "
       "does not go to zero even after calibration. For mass-level "
       "reporting this is acceptable; for safety-critical applications "
       "it would not be.")
bullet("Geometric GT uncertainty. Future work should replace the "
       "20 by 30 by 30 cm bounding-box GT with displacement-based "
       "measurements.")

# =============================================================================
# 8 Conclusion
# =============================================================================
H(1, "8. Conclusion")
P("DepthAnythingV2-Metric-Indoor-Large, combined with full-frame "
  "floor-plane fitting and a one-parameter calibration factor, "
  "estimates mixed solid-waste pile volumes from a single overhead RGB "
  "image with errors below 8 percent on all three tested conditions "
  "and below 3 percent on two of three. All 34 captures pass the "
  "centre-vs-edge depth sanity check. Raw predictions require "
  "per-condition calibration (CFs ranging 0.48 to 1.38); a single "
  "global CF is insufficient across mass levels. Calibration factor "
  "stability and per-pile CV both improve quickly with more labelled "
  "images, supporting the central empirical claim of the report: as "
  "the labelled set grows, the calibration becomes more reliable and "
  "the predictions become better.")
P("Several extensions are natural. Displacement-based GT will remove "
  "the 18 L bounding-box uncertainty. A second-order calibration "
  "(e.g.\\ linear in scene statistics such as fill fraction or mean "
  "predicted height) should reduce the residual variance after CF. "
  "Replacing the manual containment frame with a printable AprilTag "
  "scaffold would broaden the deployment envelope without sacrificing "
  "scene reference. Finally, evaluating cross-site CF transfer is the "
  "natural next step toward true field deployment.")

# =============================================================================
# Acknowledgements + References
# =============================================================================
H(1, "Acknowledgements")
P("Thanks to Praneel sir from PHO for the on-site support during data "
  "collection at the IIT Bombay solid waste site.")

H(1, "References")
P("Yang, L., Kang, B., Huang, Z., Xu, X., Feng, J., Zhao, H. "
  "Depth Anything V2. NeurIPS 2024.")
P("Kirillov, A., et al. Segment Anything. ICCV 2023.")
P("Oquab, M., et al. DINOv2: Learning Robust Visual Features without "
  "Supervision. TMLR 2023.")

doc.save(OUT_DOCX)
print(f"Saved {OUT_DOCX}")
